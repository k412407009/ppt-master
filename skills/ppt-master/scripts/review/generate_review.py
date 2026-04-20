"""
Review Board · Step 8 二次确认评审产出生成器

读 <project_dir>/review/<project>_review.json (Schema 见 references/review-board.md §VI)
产出 3 件套:
  - <project>_review.docx       完整评审报告
  - <project>_review.xlsx       问题清单 (3 sheet)
  - <project>_subjective_responses.md   主观问题最优解

Usage:
    python skills/ppt-master/scripts/review/generate_review.py <project_dir>
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


DIMENSIONS = {
    "D1": "战略-题材匹配度",
    "D2": "玩法-核心循环",
    "D3": "玩法-时间节点",
    "D4": "玩法-阶段过渡",
    "D5": "商业化-付费/留存",
    "D6": "风险-题材/合规",
    "D7": "美术/配色/素材",
}

VERDICT_LABEL = {
    "pass": "通过 (PASS)",
    "conditional_pass": "有条件通过 (CONDITIONAL PASS)",
    "not_pass": "不通过 (REJECT)",
}

PRIORITY_COLOR = {
    "P0": "FFCCCC",
    "P1": "FFE5B4",
    "P2": "DCE6F1",
}


# ============== display helpers ==============
# JSON 内部保留 P/S1/D8 等代号 (机器可读外键), 但所有面向人的产物一律展开成全名。

def _rev_label(rev: dict) -> str:
    """评委显示标签: '资深制作人 (15年)'。"""
    return f"{rev['name']} ({rev['years']}年)"


def _rev_short(rev: dict) -> str:
    """评委窄列标签 (Excel/Word 表头, 用换行分行): '资深制作人\\n15年'。"""
    return f"{rev['name']}\n{rev['years']}年"


def _rev_lookup(reviewers: list[dict]) -> dict[str, dict]:
    """{ 'P': {...}, 'S1': {...}, ... } 方便按代号反查全名。"""
    return {r["id"]: r for r in reviewers}


# ============== docx helpers ==============

def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run.font.color.rgb = RGBColor(0x1A, 0x2C, 0x5C)


def _add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold


def _add_bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(10.5)


def _scores_table(doc: Document, scores: dict[str, dict[str, int]], reviewers: list[dict]) -> None:
    n_rev = len(reviewers)
    cols = 1 + n_rev + 1  # dim + each reviewer + avg
    rows = 1 + len(DIMENSIONS) + 1  # header + dims + total

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Light Grid Accent 1"

    # header
    header = table.rows[0].cells
    header[0].text = "维度"
    _set_cell_bg(header[0], "1A2C5C")
    for i, rev in enumerate(reviewers):
        header[1 + i].text = _rev_short(rev)
        _set_cell_bg(header[1 + i], "1A2C5C")
    header[-1].text = "均分"
    _set_cell_bg(header[-1], "1A2C5C")

    for c in header:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # rows
    col_avgs: list[float] = []
    for r_idx, (dim_id, dim_name) in enumerate(DIMENSIONS.items(), start=1):
        row = table.rows[r_idx].cells
        row[0].text = dim_name
        scores_for_dim = []
        for i, rev in enumerate(reviewers):
            v = scores.get(rev["id"], {}).get(dim_id, 0)
            row[1 + i].text = str(v)
            if v:
                scores_for_dim.append(v)
        avg = round(sum(scores_for_dim) / len(scores_for_dim), 2) if scores_for_dim else 0
        row[-1].text = str(avg)
        col_avgs.append(avg)
        for c in row:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(10)

    # total row
    total = table.rows[-1].cells
    total[0].text = "评委均分"
    _set_cell_bg(total[0], "FFD700")
    for i, rev in enumerate(reviewers):
        rs = [scores.get(rev["id"], {}).get(d, 0) for d in DIMENSIONS]
        rs = [s for s in rs if s]
        avg = round(sum(rs) / len(rs), 2) if rs else 0
        total[1 + i].text = str(avg)
        _set_cell_bg(total[1 + i], "FFD700")
    overall = round(sum(col_avgs) / len(col_avgs), 2) if col_avgs else 0
    total[-1].text = str(overall)
    _set_cell_bg(total[-1], "FFD700")
    for c in total:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                run.font.bold = True
                run.font.size = Pt(10)


def _build_docx(data: dict[str, Any], out_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ---- 封面 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{data['project']}\n立项绿灯 · 评委会评审报告")
    title_run.font.name = "Microsoft YaHei"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x2C, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"\n评审日期: {data['review_date']}    "
        f"裁决: {VERDICT_LABEL.get(data['verdict'], data['verdict'])}    "
        f"加权总分: {data['weighted_score']}/5"
    )
    sub_run.font.name = "Microsoft YaHei"
    sub_run.font.size = Pt(12)

    doc.add_paragraph()

    # ---- 评委会 ----
    _add_heading(doc, "I. 评委会成员", level=1)
    rev_table = doc.add_table(rows=1 + len(data["reviewers"]), cols=3)
    rev_table.style = "Light Grid Accent 1"
    hdr = rev_table.rows[0].cells
    for i, name in enumerate(["评委", "工龄", "评审视角"]):
        hdr[i].text = name
        _set_cell_bg(hdr[i], "1A2C5C")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for i, rev in enumerate(data["reviewers"], start=1):
        cells = rev_table.rows[i].cells
        cells[0].text = rev["name"]
        cells[1].text = f"{rev['years']} 年"
        cells[2].text = rev["perspective"]
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ---- 执行摘要 ----
    _add_heading(doc, "II. 执行摘要", level=1)

    _add_heading(doc, "II.1 三大亮点", level=2)
    _add_bullet(doc, data.get("highlights", []))

    _add_heading(doc, "II.2 三大风险", level=2)
    _add_bullet(doc, data.get("risks", []))

    _add_heading(doc, "II.3 7 维度评分总表", level=2)
    _scores_table(doc, data["scores"], data["reviewers"])

    # ---- 逐评委意见 ----
    doc.add_page_break()
    _add_heading(doc, "III. 逐评委意见", level=1)
    rev_lookup = _rev_lookup(data["reviewers"])
    issues_by_rev: dict[str, list[dict]] = {r["id"]: [] for r in data["reviewers"]}
    for issue in data["issues"]:
        issues_by_rev.setdefault(issue["reviewer"], []).append(issue)

    for idx, rev in enumerate(data["reviewers"], start=1):
        _add_heading(doc, f"III.{idx} {_rev_label(rev)}", level=2)
        _add_para(doc, f"背景: {rev.get('background', '-')}", size=10)
        _add_para(doc, f"视角: {rev['perspective']}", size=10)

        rev_scores = data["scores"].get(rev["id"], {})
        # 7 维度打分拆 3 行 (每行 3 维度, 最后一行 1 维度), 全部用维度全名
        dim_items = list(DIMENSIONS.items())
        for chunk_start in range(0, len(dim_items), 3):
            chunk = dim_items[chunk_start:chunk_start + 3]
            line = "  ·  ".join(f"{name}: {rev_scores.get(did, '-')}" for did, name in chunk)
            _add_para(doc, ("打分:  " if chunk_start == 0 else "       ") + line, size=10, bold=True)

        rev_issues = issues_by_rev.get(rev["id"], [])
        if not rev_issues:
            _add_para(doc, "(无独立提问, 同意其他评委意见)", size=10)
            continue

        for q in rev_issues:
            _add_heading(doc, f"{q['id']} · [{q['priority']}] {q['question']}", level=3)
            _add_para(doc, f"维度: {DIMENSIONS.get(q['dimension'], q['dimension'])} · 类型: {'客观' if q['type'] == 'O' else '主观'} · 影响: {q.get('page', '-')}", size=10)
            if q["type"] == "O":
                _add_para(doc, f"改动建议: {q['suggestion']}", size=10)
            else:
                _add_para(doc, f"评委倾向: {q.get('subjective_position', '-')}", size=10)
                _add_para(doc, f"最优解辩护: {q.get('best_answer', '-')}", size=10)
                tps = q.get("talking_points", [])
                if tps:
                    _add_para(doc, "答辩话术:", size=10, bold=True)
                    _add_bullet(doc, tps)
        doc.add_paragraph()

    # ---- 修订总览 ----
    doc.add_page_break()
    _add_heading(doc, "IV. 修订总览 (Action Items)", level=1)
    for prio in ["P0", "P1", "P2"]:
        bucket = [q for q in data["issues"] if q["priority"] == prio and q["type"] == "O"]
        if not bucket:
            _add_heading(doc, f"IV.{prio} (无)", level=2)
            continue
        _add_heading(doc, f"IV.{prio} ({len(bucket)} 项)", level=2)
        for q in bucket:
            rev_name = rev_lookup[q["reviewer"]]["name"] if q["reviewer"] in rev_lookup else q["reviewer"]
            _add_para(
                doc,
                f"[{q['id']}] @ {q.get('page','-')} — {q['question']} → {q['suggestion']} (评委: {rev_name})",
                size=10,
            )

    # ---- 最终裁决 ----
    doc.add_page_break()
    _add_heading(doc, "V. 最终裁决", level=1)
    verdict_text = VERDICT_LABEL.get(data["verdict"], data["verdict"])
    _add_para(doc, f"裁决: {verdict_text}", bold=True, size=14)
    _add_para(doc, f"加权总分: {data['weighted_score']}/5", size=12)
    if data.get("verdict_rationale"):
        _add_para(doc, f"理由: {data['verdict_rationale']}", size=11)
    if data.get("next_review"):
        _add_para(doc, f"复审节点: {data['next_review']}", size=11)

    doc.save(out_path)


# ============== xlsx ==============

THICK_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


def _build_xlsx(data: dict[str, Any], out_path: Path) -> None:
    wb = Workbook()

    rev_lookup = _rev_lookup(data["reviewers"])

    # ---- Sheet 1: Issues ----
    ws = wb.active
    ws.title = "Issues"
    headers = ["ID", "评委", "维度", "类型", "优先级", "页号/影响", "问题", "改动建议/最优解", "答辩话术"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF", name="Microsoft YaHei")
        cell.fill = PatternFill("solid", fgColor="1A2C5C")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THICK_BORDER
    for r, q in enumerate(data["issues"], start=2):
        if q["type"] == "O":
            advice = q.get("suggestion", "-")
            tps = "-"
        else:
            advice = (
                f"[评委倾向] {q.get('subjective_position', '-')}\n"
                f"[最优解] {q.get('best_answer', '-')}"
            )
            tps = "\n".join(f"• {t}" for t in q.get("talking_points", []))
        rev = rev_lookup.get(q["reviewer"], {})
        row_vals = [
            q["id"],
            _rev_label(rev) if rev else q["reviewer"],
            DIMENSIONS.get(q["dimension"], q["dimension"]),
            "客观" if q["type"] == "O" else "主观",
            q["priority"],
            q.get("page", "-"),
            q["question"],
            advice,
            tps,
        ]
        for c, v in enumerate(row_vals, start=1):
            cell = ws.cell(row=r, column=c, value=v)
            cell.font = Font(name="Microsoft YaHei", size=10)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = THICK_BORDER
        prio_color = PRIORITY_COLOR.get(q["priority"], "FFFFFF")
        ws.cell(row=r, column=5).fill = PatternFill("solid", fgColor=prio_color)
    # 列宽: ID / 评委 / 维度 / 类型 / 优先级 / 页号/影响 / 问题 / 改动建议 / 答辩话术
    widths = [8, 22, 22, 6, 8, 16, 50, 60, 40]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    # ---- Sheet 2: Scores ----
    ws2 = wb.create_sheet("Scores")
    rev_ids = [r["id"] for r in data["reviewers"]]
    rev_col_labels = [_rev_short(r) for r in data["reviewers"]]  # 表头展开成全名+工龄
    headers2 = ["维度"] + rev_col_labels + ["均分"]
    for c, h in enumerate(headers2, start=1):
        cell = ws2.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF", name="Microsoft YaHei")
        cell.fill = PatternFill("solid", fgColor="1A2C5C")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = THICK_BORDER
    ws2.row_dimensions[1].height = 36  # 表头双行高度

    col_sums: dict[str, list[int]] = {rid: [] for rid in rev_ids}
    for r, (dim_id, dim_name) in enumerate(DIMENSIONS.items(), start=2):
        ws2.cell(row=r, column=1, value=dim_name).font = Font(bold=True, name="Microsoft YaHei")
        scores_row = []
        for c, rid in enumerate(rev_ids, start=2):
            v = data["scores"].get(rid, {}).get(dim_id, 0)
            cell = ws2.cell(row=r, column=c, value=v)
            cell.alignment = Alignment(horizontal="center")
            cell.border = THICK_BORDER
            if v:
                scores_row.append(v)
                col_sums[rid].append(v)
        avg = round(sum(scores_row) / len(scores_row), 2) if scores_row else 0
        cell = ws2.cell(row=r, column=len(headers2), value=avg)
        cell.font = Font(bold=True, name="Microsoft YaHei")
        cell.alignment = Alignment(horizontal="center")
        cell.fill = PatternFill("solid", fgColor="FFF7E6")

    # 评委均分行
    total_row = 2 + len(DIMENSIONS)
    cell = ws2.cell(row=total_row, column=1, value="评委均分")
    cell.font = Font(bold=True, name="Microsoft YaHei")
    cell.fill = PatternFill("solid", fgColor="FFD700")
    for c, rid in enumerate(rev_ids, start=2):
        rs = col_sums[rid]
        avg = round(sum(rs) / len(rs), 2) if rs else 0
        cell = ws2.cell(row=total_row, column=c, value=avg)
        cell.font = Font(bold=True, name="Microsoft YaHei")
        cell.alignment = Alignment(horizontal="center")
        cell.fill = PatternFill("solid", fgColor="FFD700")
    cell = ws2.cell(row=total_row, column=len(headers2), value=data["weighted_score"])
    cell.font = Font(bold=True, name="Microsoft YaHei")
    cell.alignment = Alignment(horizontal="center")
    cell.fill = PatternFill("solid", fgColor="FFD700")

    ws2.column_dimensions["A"].width = 26
    for c in range(2, len(headers2) + 1):
        ws2.column_dimensions[get_column_letter(c)].width = 16
    ws2.freeze_panes = "B2"

    # ---- Sheet 3: Action_Items ----
    ws3 = wb.create_sheet("Action_Items")
    a_headers = ["优先级", "ID", "页号", "问题", "改动建议", "评委", "维度", "预估工时", "Owner"]
    for c, h in enumerate(a_headers, start=1):
        cell = ws3.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF", name="Microsoft YaHei")
        cell.fill = PatternFill("solid", fgColor="1A2C5C")
        cell.alignment = Alignment(horizontal="center")
        cell.border = THICK_BORDER

    rownum = 2
    for prio in ["P0", "P1", "P2"]:
        bucket = [q for q in data["issues"] if q["priority"] == prio and q["type"] == "O"]
        for q in bucket:
            rev = rev_lookup.get(q["reviewer"], {})
            row_vals = [
                prio,
                q["id"],
                q.get("page", "-"),
                q["question"],
                q.get("suggestion", "-"),
                _rev_label(rev) if rev else q["reviewer"],
                DIMENSIONS.get(q["dimension"], q["dimension"]),
                q.get("est_hours", ""),
                q.get("owner", ""),
            ]
            for c, v in enumerate(row_vals, start=1):
                cell = ws3.cell(row=rownum, column=c, value=v)
                cell.font = Font(name="Microsoft YaHei", size=10)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = THICK_BORDER
            ws3.cell(row=rownum, column=1).fill = PatternFill("solid", fgColor=PRIORITY_COLOR[prio])
            rownum += 1
    # 列宽: 优先级 / ID / 页号 / 问题 / 改动建议 / 评委 / 维度 / 预估工时 / Owner
    a_widths = [8, 8, 16, 50, 60, 22, 22, 10, 14]
    for i, w in enumerate(a_widths, start=1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    wb.save(out_path)


# ============== subjective md ==============

def _build_subjective_md(data: dict[str, Any], out_path: Path) -> None:
    rev_lookup = _rev_lookup(data["reviewers"])
    lines = [
        f"# {data['project']} · 主观问题最优解 (Subjective Responses)",
        "",
        f"> 评委会 = {' / '.join(_rev_label(r) for r in data['reviewers'])}",
        f"> 评审日期: {data['review_date']}    裁决: {VERDICT_LABEL.get(data['verdict'], data['verdict'])}",
        "",
        "本文档仅含主观/无定论问题, 提供评委倾向 + PPT 当前方案的最优解辩护 + 答辩话术。",
        "客观问题 + 改动建议见同目录 `*_review.docx` 第 IV 章 / `*_review.xlsx` Action_Items sheet。",
        "",
        "---",
        "",
    ]
    s_issues = [q for q in data["issues"] if q["type"] == "S"]
    if not s_issues:
        lines.append("(本次评审无主观问题, 全部为可执行客观项)")
    else:
        for q in s_issues:
            rev = rev_lookup.get(q["reviewer"], {})
            rev_str = _rev_label(rev) if rev else q["reviewer"]
            lines.extend([
                f"## {q['id']} · {q['question']}",
                "",
                f"- **评委**: {rev_str} · **维度**: {DIMENSIONS.get(q['dimension'], q['dimension'])} · **优先级**: {q['priority']} · **影响**: {q.get('page', '-')}",
                "",
                f"**评委倾向**:  {q.get('subjective_position', '-')}",
                "",
                f"**最优解辩护**:  {q.get('best_answer', '-')}",
                "",
                "**答辩话术**:",
                "",
            ])
            tps = q.get("talking_points", [])
            if tps:
                for t in tps:
                    lines.append(f"- {t}")
            else:
                lines.append("- (评委此问无明确话术建议, 制作人临场发挥)")
            lines.extend(["", "---", ""])
    out_path.write_text("\n".join(lines), encoding="utf-8")


# ============== main ==============

def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("用法: python generate_review.py <project_dir>")
        return 2
    project_dir = Path(argv[1]).resolve()
    if not project_dir.exists():
        print(f"项目目录不存在: {project_dir}")
        return 2

    review_dir = project_dir / "review"
    review_dir.mkdir(exist_ok=True)

    json_files = sorted(review_dir.glob("*_review.json"))
    if not json_files:
        print(f"未找到 review json: {review_dir}/*_review.json")
        return 2
    json_path = json_files[0]
    data = json.loads(json_path.read_text(encoding="utf-8"))

    project_slug = data["project"].replace(" ", "_").replace("·", "_").replace("/", "_")
    docx_path = review_dir / f"{project_slug}_review.docx"
    xlsx_path = review_dir / f"{project_slug}_review.xlsx"
    md_path = review_dir / f"{project_slug}_subjective_responses.md"

    print(f"  reading: {json_path}")
    _build_docx(data, docx_path)
    print(f"  wrote  : {docx_path}")
    _build_xlsx(data, xlsx_path)
    print(f"  wrote  : {xlsx_path}")
    _build_subjective_md(data, md_path)
    print(f"  wrote  : {md_path}")

    print(f"评审三件套已生成 → {review_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
